#!/usr/bin/env python3
"""
舍得日报生成核心模块
"""
import sys
import json
import shutil
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile

TEMPLATE = os.path.join(os.path.dirname(__file__), '..', 'skills', 'sheda-daily-report', 'assets', 'template_with_logo.docx')
LOGO = os.path.join(os.path.dirname(__file__), '..', 'skills', 'sheda-daily-report', 'assets', 'logo.png')

def week_day(date_str):
    from datetime import datetime
    dt = datetime.strptime(date_str, '%Y-%m-%d')
    return ['星期日','星期一','星期二','星期三','星期四','星期五','星期六'][dt.weekday()]

def make_report(date_str, weather, content, photo_paths, output_path):
    date_display = f"{date_str[:4]}年{int(date_str[5:7])}月{int(date_str[8:])}日"
    wd = week_day(date_str)

    # 复制模板
    if os.path.exists(TEMPLATE):
        shutil.copy(TEMPLATE, output_path)
        doc = Document(output_path)
    else:
        doc = Document()

    body = doc.element.body
    for p in list(doc.paragraphs): p._element.getparent().remove(p._element)
    for pic in body.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic'): pic.getparent().remove(pic)
    for tbl in body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'): tbl.getparent().remove(tbl)

    def add_para(text, font_name='宋体', font_size=12, bold=False,
                 align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.25):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.line_spacing = line_spacing
        pf.first_line_indent = 0
        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        return p

    # 主标题
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.line_spacing = Pt(26)
    r0 = p0.add_run('工程日报')
    r0.font.name = '宋体'
    r0.font.size = Pt(22)
    r0.bold = True
    r0._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_para(f'日期：{date_display} {wd} 天气：{weather}')
    add_para('项目名称：舍得酒业酒文化博物馆建设项目夯土墙施工日报')
    doc.add_paragraph('')

    sections = [
        ('一、基本情况', content.get('基本情况', [])),
        ('二、人员设备', content.get('人员设备', [])),
        ('三、当日完成', content.get('当日完成', [])),
        ('四、次日计划', content.get('次日计划', [])),
        ('五、需要协调解决问题', [content.get('协调问题', '无')]),
    ]
    for title, lines in sections:
        add_para(title)
        for line in lines:
            add_para(line)
        doc.add_paragraph('')

    # 照片
    if photo_paths:
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_pr = table._element.find('./{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            table._element.insert(0, tbl_pr)
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '9588')
        tbl_w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')
        tbl_pr.append(tbl_w)

        widths = ['4740', '4848']
        for ci, photo_path in enumerate(photo_paths[:2]):
            cell = table.cell(0, ci)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement('w:tcW')
            tc_w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', widths[ci])
            tc_w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')
            tc_pr.append(tc_w)
            for p in list(cell.paragraphs): p._element.getparent().remove(p._element)
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(photo_path, width=Cm(7.8))

        captions = ['三级教育培训'] * len(photo_paths)
        for ci, caption in enumerate(captions[:2]):
            cell = table.cell(1, ci)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement('w:tcW')
            tc_w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', widths[ci])
            tc_w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')
            tc_pr.append(tc_w)
            for p in list(cell.paragraphs): p._element.getparent().remove(p._element)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            r.font.name = '宋体'
            r.font.size = Pt(10.5)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.save(output_path)


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--date', required=True)
    parser.add_argument('--weather', required=True)
    parser.add_argument('--content', required=True)
    parser.add_argument('--output', required=True)
    parser.add_argument('--photos', nargs='*', default=[])
    args = parser.parse_args()

    content = json.loads(args.content)
    make_report(args.date, args.weather, content, args.photos, args.output)
    print('OK:' + args.output)
